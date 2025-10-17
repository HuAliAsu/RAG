"""
مدیریت خواندن و نوشتن فایل‌های Word با metadata کامل
"""

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from datetime import datetime
import hashlib


class WordDocumentHandler:
    """
    کلاس مدیریت اسناد Word
    """

    def __init__(self):
        """مقداردهی اولیه"""
        pass

    def read_document(self, file_path: str) -> Tuple[str, List[Dict], Dict]:
        """
        خواندن فایل Word و استخراج متن + ساختار

        Args:
            file_path: مسیر فایل

        Returns:
            tuple: (full_text, headings, metadata)
        """
        from src.utils.logger import info, debug

        doc = Document(file_path)

        full_text_parts = []
        headings = []
        paragraphs_info = []

        para_index = 0

        for para in doc.paragraphs:
            text = para.text.strip()

            # اطلاعات پاراگراف
            para_info = {
                'index': para_index,
                'text': text,
                'style': para.style.name,
                'is_heading': False,
                'heading_level': 0
            }

            if not text:
                full_text_parts.append('')
                para_index += 1
                paragraphs_info.append(para_info)
                continue

            # بررسی Heading
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                level = 1
                if 'heading 1' in style_name:
                    level = 1
                elif 'heading 2' in style_name:
                    level = 2
                elif 'heading 3' in style_name:
                    level = 3

                headings.append({
                    'text': text,
                    'level': level,
                    'para_index': para_index
                })

                para_info['is_heading'] = True
                para_info['heading_level'] = level

                debug(f"   Heading {level} یافت شد: {text[:50]}")

            full_text_parts.append(text)
            paragraphs_info.append(para_info)
            para_index += 1

        full_text = '\n'.join(full_text_parts)

        # ایجاد metadata سند
        file_obj = Path(file_path)
        metadata = {
            'source_file': file_obj.name,
            'file_path': str(file_obj.absolute()),
            'file_size': file_obj.stat().st_size,
            'total_paragraphs': para_index,
            'total_headings': len(headings),
            'extracted_at': datetime.now().isoformat(),
            'file_hash': self._calculate_file_hash(file_path)
        }

        info(f"✅ سند خوانده شد: {metadata['total_paragraphs']} پاراگراف، {len(headings)} عنوان")

        return full_text, headings, metadata

    def _calculate_file_hash(self, file_path: str) -> str:
        """محاسبه hash فایل"""
        hasher = hashlib.sha256()
        with open(file_path, 'rb') as f:
            hasher.update(f.read())
        return hasher.hexdigest()[:16]

    def save_chunked_document(self, output_path: str, chunks: List[Dict],
                              source_metadata: Dict, processing_stats: Dict):
        """
        ذخیره فایل Word چانک‌شده با تگ‌ها و metadata

        Args:
            output_path: مسیر خروجی
            chunks: لیست chunk‌ها
            source_metadata: metadata فایل اصلی
            processing_stats: آمار پردازش
        """
        from src.utils.logger import info

        doc = Document()

        # عنوان سند
        title = doc.add_heading('سند چانک‌شده - RAG Workbench', 0)

        # اطلاعات کلی
        info_para = doc.add_paragraph()
        info_para.add_run('فایل اصلی: ').bold = True
        info_para.add_run(f"{source_metadata['source_file']}\n")
        info_para.add_run('تاریخ پردازش: ').bold = True
        info_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run('تعداد Chunk: ').bold = True
        info_para.add_run(f"{processing_stats['total_chunks']}\n")
        info_para.add_run('زمان پردازش: ').bold = True
        info_para.add_run(f"{processing_stats['total_time']:.2f} ثانیه\n")

        doc.add_paragraph()
        doc.add_paragraph('─' * 80)
        doc.add_paragraph()

        # نوشتن chunk‌ها
        for chunk in chunks:
            # تگ chunk
            tag = f"@@-{chunk['chunk_id']}-@@"
            tag_para = doc.add_paragraph()
            tag_run = tag_para.add_run(tag)
            tag_run.bold = True
            tag_run.font.size = Pt(14)
            tag_run.font.color.rgb = RGBColor(0, 102, 204)  # آبی

            # نوع chunk (Parent/Child)
            if chunk.get('is_parent'):
                type_para = doc.add_paragraph()
                type_run = type_para.add_run('📦 Parent Chunk')
                type_run.font.color.rgb = RGBColor(0, 153, 76)  # سبز
            else:
                type_para = doc.add_paragraph()
                type_run = type_para.add_run('📄 Child Chunk')
                type_run.font.color.rgb = RGBColor(204, 102, 0)  # نارنجی

            # Heading path (اگر دارد)
            if chunk.get('heading_path'):
                heading_para = doc.add_paragraph()
                heading_para.add_run('📍 مسیر: ').bold = True
                heading_para.add_run(' » '.join(chunk['heading_path']))

            # متن chunk با highlight
            text_para = doc.add_paragraph()
            text_run = text_para.add_run(chunk['text'])

            # highlight بر اساس boundary stage
            stage = chunk['boundary_info'].get('stage', 0)
            if stage == 0:  # Heading
                text_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif stage == 1:  # Pattern
                text_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            elif stage == 2:  # Coherence
                text_run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            elif stage == 3:  # Embedding
                text_run.font.highlight_color = WD_COLOR_INDEX.PINK

            # Metadata
            meta_para = doc.add_paragraph()
            meta_para.style = 'Intense Quote'

            meta_parts = [
                f"ID: {chunk['chunk_id']}",
                f"کلمات: {chunk['word_count']}",
                f"جملات: {chunk['sentence_count']}",
                f"Stage: {chunk['boundary_info'].get('stage_name', 'N/A')}",
                f"اطمینان: {chunk['boundary_info'].get('confidence', 0):.2f}"
            ]

            if chunk.get('parent_id'):
                meta_parts.append(f"والد: {chunk['parent_id']}")

            if chunk.get('child_ids'):
                meta_parts.append(f"فرزندان: {len(chunk['child_ids'])}")

            if chunk.get('keywords'):
                meta_parts.append(f"کلیدواژه: {', '.join(chunk['keywords'][:3])}")

            meta_para.add_run(' | '.join(meta_parts))

            # فاصله
            doc.add_paragraph()
            doc.add_paragraph('─' * 80)
            doc.add_paragraph()

        # ذخیره
        doc.save(output_path)
        info(f"✅ فایل چانک‌شده ذخیره شد: {output_path}")

    def extract_paragraph_ids(self, text: str) -> List[int]:
        """
        استخراج شماره پاراگراف‌ها از متن
        (برای ساخت ارتباط بین chunk و پاراگراف اصلی)
        """
        # این متد در فازهای بعدی تکمیل می‌شود
        return []


# تست
if __name__ == "__main__":
    print("🧪 تست word_handler.py\n")

    handler = WordDocumentHandler()

    # تست خواندن (اگر فایل تست وجود دارد)
    test_file = "./Word_File_for_Test/Word_File_for_Test.docx"
    if Path(test_file).exists():
        text, headings, metadata = handler.read_document(test_file)
        print(f"✅ متن خوانده شد: {len(text)} کاراکتر")
        print(f"✅ {len(headings)} عنوان یافت شد")
        print(f"✅ Metadata: {metadata}")
    else:
        print(f"⚠️ فایل تست یافت نشد: {test_file}")