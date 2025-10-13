"""
سیستم Chunking هوشمند سه مرحله‌ای برای متن‌های فارسی
Optimized Hybrid: Pattern + Coherence + Semantic Embedding
نویسنده: کد تولید شده برای RAG
"""

import re
import numpy as np
from typing import List, Dict, Tuple, Optional
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import hazm
from pathlib import Path
import requests
import json
from collections import Counter
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


class OllamaEmbedding:
    """کلاس برای تعامل با Ollama"""

    def __init__(self, model: str = "embeddinggemma", base_url: str = "http://localhost:11434"):
        self.model = model
        self.base_url = base_url
        self.embed_url = f"{base_url}/api/embeddings"

    def get_embedding(self, text: str) -> Optional[np.ndarray]:
        """دریافت embedding از Ollama"""
        try:
            response = requests.post(
                self.embed_url,
                json={"model": self.model, "prompt": text},
                timeout=30
            )

            if response.status_code == 200:
                result = response.json()
                return np.array(result['embedding'])
            else:
                print(f"⚠️ خطا در دریافت embedding: {response.status_code}")
                return None

        except requests.exceptions.RequestException as e:
            print(f"⚠️ خطا در اتصال به Ollama: {e}")
            return None

    def get_batch_embeddings(self, texts: List[str], show_progress: bool = True) -> List[Optional[np.ndarray]]:
        """دریافت embedding برای لیست متن‌ها"""
        embeddings = []
        total = len(texts)

        for idx, text in enumerate(texts):
            if show_progress and (idx % 10 == 0 or idx == total - 1):
                print(f"  Embedding: {idx + 1}/{total}", end='\r')

            emb = self.get_embedding(text)
            embeddings.append(emb)

        if show_progress:
            print()  # خط جدید بعد از progress

        return embeddings


class OptimizedHybridChunker:
    """Chunker سه مرحله‌ای بهینه شده"""

    def __init__(
            self,
            min_chunk_size: int = 200,
            max_chunk_size: int = 800,
            overlap_size: int = 50,
            ollama_model: str = "embeddinggemma"
    ):
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size

        # ابزارهای پردازش متن
        self.normalizer = hazm.Normalizer()
        self.sent_tokenizer = hazm.SentenceTokenizer()
        self.word_tokenizer = hazm.WordTokenizer()

        # Ollama embedding
        self.embedder = OllamaEmbedding(model=ollama_model)

        # آمار
        self.stats = {
            'stage1_boundaries': 0,
            'stage2_boundaries': 0,
            'stage3_boundaries': 0,
            'embedding_calls': 0
        }

    def normalize_text(self, text: str) -> str:
        """نرمال‌سازی متن فارسی"""
        text = self.normalizer.normalize(text)
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    # ============================================
    # STAGE 1: Fast Pattern-Based Detection
    # ============================================

    def stage1_pattern_detection(self, paragraphs: List[str]) -> List[Dict]:
        """مرحله 1: تشخیص سریع با الگوها"""
        boundaries = []

        for idx, para in enumerate(paragraphs):
            if not para.strip():
                continue

            score = 0.0
            confidence = 0.0
            signals = []

            # الگوهای قوی ساختاری (confidence: 1.0)
            structural_patterns = [
                (r'^(فصل|بخش)\s+[\d۰-۹]+', 1.0, 'chapter'),
                (r'^داستان\s+[\d۰-۹]+', 1.0, 'story_title'),
                (r'^[\d۰-۹]+[\.\)]\s+', 0.9, 'numbered'),
                (r'^={3,}|^-{3,}|^\*{3,}', 0.95, 'separator'),
                (r'^(اول|دوم|سوم|چهارم|پنجم|ششم|هفتم|هشتم|نهم|دهم)\s*[-:]\s*', 0.9, 'ordinal'),
            ]

            for pattern, conf, signal in structural_patterns:
                if re.search(pattern, para):
                    score += 3.0
                    confidence = max(confidence, conf)
                    signals.append(signal)
                    break

            # الگوهای معنایی متوسط (confidence: 0.6-0.8)
            semantic_patterns = [
                # شروع داستان
                (r'(روزی\s+روزگاری|یکی\s+بود\s+یکی\s+نبود)', 0.7, 'story_start'),
                (r'در\s+(زمانی|روزی|شهری|دیاری)\s+که', 0.65, 'context_start'),

                # پایان بخش
                (r'(درس|نکته|آموخته|عبرت)[:\s]+', 0.75, 'lesson'),
                (r'به\s+این\s+ترتیب', 0.6, 'conclusion'),

                # نشانه‌های گفتمان
                (r'^(حال|اکنون|بگذارید)\s+', 0.65, 'discourse_marker'),
                (r'^(اما|ولی|با\s+این\s+حال)\s+', 0.6, 'contrast'),
            ]

            for pattern, conf, signal in semantic_patterns:
                if re.search(pattern, para):
                    score += 2.0
                    confidence = max(confidence, conf)
                    signals.append(signal)

            # الگوهای ساختار پاراگراف (confidence: 0.5-0.7)
            words = self.word_tokenizer.tokenize(para)
            para_length = len(words)

            # پاراگراف کوتاه + موقعیت مناسب
            if para_length < 15 and idx > 0:
                prev_para = paragraphs[idx - 1] if idx > 0 else ""
                next_para = paragraphs[idx + 1] if idx < len(paragraphs) - 1 else ""

                if len(prev_para.strip()) > 50 and len(next_para.strip()) > 50:
                    score += 1.5
                    confidence = max(confidence, 0.7)
                    signals.append('short_title')

            # خطوط خالی قبل و بعد
            has_empty_before = idx > 0 and not paragraphs[idx - 1].strip()
            has_empty_after = idx < len(paragraphs) - 1 and not paragraphs[idx + 1].strip()

            if has_empty_before and has_empty_after and para_length < 20:
                score += 1.5
                confidence = max(confidence, 0.65)
                signals.append('isolated')

            # پاراگراف خیلی بلند (احتمال تغییر موضوع)
            if para_length > 200:
                score += 0.5
                confidence = max(confidence, 0.4)
                signals.append('long_para')

            # ذخیره مرز
            if score > 0:
                boundaries.append({
                    'index': idx,
                    'score': score,
                    'confidence': confidence,
                    'text': para,
                    'signals': signals,
                    'stage': 1
                })

        # فیلتر بر اساس confidence
        high_conf = [b for b in boundaries if b['confidence'] >= 0.8]
        medium_conf = [b for b in boundaries if 0.5 <= b['confidence'] < 0.8]
        low_conf = [b for b in boundaries if b['confidence'] < 0.5]

        self.stats['stage1_boundaries'] = len(high_conf)

        print(f"  Stage 1: {len(high_conf)} مرز قطعی | {len(medium_conf)} مرز احتمالی | {len(low_conf)} مرز ضعیف")

        return {
            'confirmed': high_conf,
            'uncertain': medium_conf + low_conf,
            'all': boundaries
        }

    # ============================================
    # STAGE 2: Coherence Analysis
    # ============================================

    def stage2_coherence_analysis(self, paragraphs: List[str], uncertain_boundaries: List[Dict]) -> List[Dict]:
        """مرحله 2: تحلیل انسجام متن"""

        if not uncertain_boundaries:
            return []

        confirmed = []
        still_uncertain = []

        for boundary in uncertain_boundaries:
            idx = boundary['index']

            # محاسبه coherence قبل و بعد از مرز
            before_coherence = self._calculate_coherence(paragraphs, max(0, idx - 3), idx)
            after_coherence = self._calculate_coherence(paragraphs, idx, min(len(paragraphs), idx + 3))
            cross_coherence = self._calculate_coherence(paragraphs, max(0, idx - 2), min(len(paragraphs), idx + 2))

            # اگر coherence درون بخش‌ها بالا و بین بخش‌ها پایین → مرز قطعی
            coherence_drop = (before_coherence + after_coherence) / 2 - cross_coherence

            if coherence_drop > 0.15:  # آستانه تجربی
                boundary['confidence'] = 0.8
                boundary['coherence_drop'] = coherence_drop
                boundary['stage'] = 2
                confirmed.append(boundary)
            elif coherence_drop > 0.08:
                boundary['coherence_drop'] = coherence_drop
                still_uncertain.append(boundary)
            # else: رد می‌شود

        self.stats['stage2_boundaries'] = len(confirmed)

        print(f"  Stage 2: {len(confirmed)} مرز تأیید شد | {len(still_uncertain)} مرز نیاز به embedding")

        return {
            'confirmed': confirmed,
            'uncertain': still_uncertain
        }

    def _calculate_coherence(self, paragraphs: List[str], start: int, end: int) -> float:
        """محاسبه انسجام متنی بین پاراگراف‌ها"""

        if start >= end or start < 0 or end > len(paragraphs):
            return 0.0

        section = ' '.join([p for p in paragraphs[start:end] if p.strip()])

        if not section:
            return 0.0

        # استخراج کلمات کلیدی
        words = self.word_tokenizer.tokenize(section)

        # حذف stop words
        stop_words = {'و', 'در', 'به', 'از', 'که', 'این', 'را', 'با', 'برای', 'آن', 'یک', 'است', 'شود', 'می'}
        filtered_words = [w for w in words if len(w) > 2 and w not in stop_words]

        if len(filtered_words) < 5:
            return 0.0

        # محاسبه تکرار کلمات (نشانه انسجام)
        word_freq = Counter(filtered_words)
        repeated_words = sum(1 for count in word_freq.values() if count > 1)

        # نرمال‌سازی
        coherence = repeated_words / len(set(filtered_words)) if filtered_words else 0

        return min(coherence, 1.0)

    # ============================================
    # STAGE 3: Semantic Embedding
    # ============================================

    def stage3_semantic_verification(self, paragraphs: List[str], uncertain_boundaries: List[Dict]) -> List[Dict]:
        """مرحله 3: تأیید با embedding معنایی"""

        if not uncertain_boundaries:
            print("  Stage 3: هیچ مرز نامشخصی برای بررسی نیست")
            return []

        print(f"  Stage 3: بررسی {len(uncertain_boundaries)} مرز با embedding...")

        confirmed = []

        # برای هر مرز نامشخص
        for boundary in uncertain_boundaries:
            idx = boundary['index']

            # متن قبل و بعد از مرز
            before_text = ' '.join([p for p in paragraphs[max(0, idx - 2):idx] if p.strip()])
            after_text = ' '.join([p for p in paragraphs[idx:min(len(paragraphs), idx + 3)] if p.strip()])

            if not before_text or not after_text:
                continue

            # دریافت embedding
            emb_before = self.embedder.get_embedding(before_text)
            emb_after = self.embedder.get_embedding(after_text)

            self.stats['embedding_calls'] += 2

            if emb_before is None or emb_after is None:
                print(f"    ⚠️ خطا در embedding برای مرز {idx}")
                continue

            # محاسبه شباهت
            similarity = np.dot(emb_before, emb_after) / (np.linalg.norm(emb_before) * np.linalg.norm(emb_after))

            # اگر شباهت پایین → مرز قطعی
            if similarity < 0.75:  # آستانه تجربی
                boundary['confidence'] = 0.85
                boundary['semantic_similarity'] = float(similarity)
                boundary['stage'] = 3
                confirmed.append(boundary)

        self.stats['stage3_boundaries'] = len(confirmed)

        print(f"  Stage 3: {len(confirmed)} مرز تأیید شد با embedding")

        return confirmed

    # ============================================
    # Main Processing Pipeline
    # ============================================

    def create_chunks(self, text: str, headings: List[Dict] = None) -> List[Dict]:
        """پایپلاین اصلی chunking سه مرحله‌ای"""

        print("\n🔄 شروع Chunking سه مرحله‌ای...")

        text = self.normalize_text(text)
        paragraphs = text.split('\n')

        # افزودن heading‌ها به عنوان مرزهای قطعی
        confirmed_boundaries = []

        if headings:
            print(f"  📑 {len(headings)} عنوان (Heading) تشخیص داده شد")
            for heading in headings:
                confirmed_boundaries.append({
                    'index': heading['para_index'],
                    'score': 5.0,
                    'confidence': 1.0,
                    'text': heading['text'],
                    'signals': ['heading'],
                    'stage': 0,  # از فایل ورد
                    'heading_level': heading['level']
                })

        # Stage 1: Pattern Detection
        print("\n⚡ Stage 1: Pattern Detection (سریع)")
        stage1_result = self.stage1_pattern_detection(paragraphs)
        confirmed_boundaries.extend(stage1_result['confirmed'])

        # Stage 2: Coherence Analysis
        print("\n🔍 Stage 2: Coherence Analysis (متوسط)")
        stage2_result = self.stage2_coherence_analysis(paragraphs, stage1_result['uncertain'])
        confirmed_boundaries.extend(stage2_result['confirmed'])

        # Stage 3: Semantic Embedding
        print("\n🧠 Stage 3: Semantic Embedding (دقیق)")
        stage3_result = self.stage3_semantic_verification(paragraphs, stage2_result['uncertain'])
        confirmed_boundaries.extend(stage3_result)

        # مرتب‌سازی بر اساس index
        confirmed_boundaries.sort(key=lambda x: x['index'])

        # ایجاد chunk‌ها
        print("\n✂️ ایجاد chunk‌ها...")
        chunks = self._create_chunks_from_boundaries(paragraphs, confirmed_boundaries)

        # افزودن overlap
        chunks = self._add_smart_overlap(chunks)

        # بهینه‌سازی نهایی
        chunks = self._optimize_chunks(chunks)

        # آمار نهایی
        self._print_stats(len(chunks))

        return chunks

    def _create_chunks_from_boundaries(self, paragraphs: List[str], boundaries: List[Dict]) -> List[Dict]:
        """ایجاد chunk‌ها از مرزهای تشخیص داده شده"""

        chunks = []
        boundary_indices = [b['index'] for b in boundaries]
        boundary_indices = [0] + boundary_indices + [len(paragraphs)]

        for i in range(len(boundary_indices) - 1):
            start_idx = boundary_indices[i]
            end_idx = boundary_indices[i + 1]

            section = '\n'.join([p for p in paragraphs[start_idx:end_idx] if p.strip()])

            if not section.strip():
                continue

            # metadata مرز
            boundary_info = None
            if i < len(boundaries):
                boundary_info = boundaries[i]

            # ایجاد chunk
            chunk = self._create_chunk_dict(section, len(chunks), boundary_info)
            chunks.append(chunk)

        return chunks

    def _create_chunk_dict(self, text: str, chunk_id: int, boundary_info: Optional[Dict]) -> Dict:
        """ایجاد dictionary برای chunk"""

        words = self.word_tokenizer.tokenize(text)
        sentences = self.sent_tokenizer.tokenize(text)

        chunk = {
            'chunk_id': chunk_id,
            'text': text,
            'word_count': len(words),
            'sentence_count': len(sentences),
            'keywords': self._extract_keywords(text),
            'overlap_prev': '',
            'overlap_next': ''
        }

        # اضافه کردن اطلاعات مرز
        if boundary_info:
            chunk['boundary_confidence'] = boundary_info.get('confidence', 0)
            chunk['boundary_stage'] = boundary_info.get('stage', 0)
            chunk['boundary_signals'] = boundary_info.get('signals', [])

            if 'heading_level' in boundary_info:
                chunk['heading_level'] = boundary_info['heading_level']

        return chunk

    def _add_smart_overlap(self, chunks: List[Dict]) -> List[Dict]:
        """افزودن overlap هوشمند"""

        for i in range(len(chunks) - 1):
            curr_chunk = chunks[i]
            next_chunk = chunks[i + 1]

            # استخراج جملات آخر
            curr_sents = self.sent_tokenizer.tokenize(curr_chunk['text'])

            # تعداد جملات overlap بسته به اطمینان مرز
            confidence = next_chunk.get('boundary_confidence', 0.5)

            if confidence >= 0.9:  # مرز قطعی → overlap کم
                overlap_count = 1
            elif confidence >= 0.7:  # مرز متوسط → overlap متوسط
                overlap_count = 2
            else:  # مرز ضعیف → overlap زیاد
                overlap_count = 3

            overlap_sents = curr_sents[-overlap_count:] if len(curr_sents) >= overlap_count else curr_sents
            overlap_text = ' '.join(overlap_sents)

            curr_chunk['overlap_next'] = overlap_text
            next_chunk['overlap_prev'] = overlap_text

        return chunks

    def _optimize_chunks(self, chunks: List[Dict]) -> List[Dict]:
        """بهینه‌سازی نهایی chunk‌ها"""

        optimized = []
        i = 0

        while i < len(chunks):
            chunk = chunks[i]

            # ادغام chunk‌های خیلی کوچک
            if chunk['word_count'] < self.min_chunk_size and i < len(chunks) - 1:
                next_chunk = chunks[i + 1]

                # بررسی اینکه آیا ادغام منطقی است
                if next_chunk['word_count'] < self.max_chunk_size:
                    merged_text = chunk['text'] + '\n\n' + next_chunk['text']
                    chunk['text'] = merged_text
                    chunk['word_count'] = len(self.word_tokenizer.tokenize(merged_text))
                    chunk['merged'] = True
                    optimized.append(chunk)
                    i += 2
                    continue

            # تقسیم chunk‌های خیلی بزرگ
            if chunk['word_count'] > self.max_chunk_size:
                sub_chunks = self._split_large_chunk(chunk['text'], chunk['chunk_id'])
                optimized.extend(sub_chunks)
            else:
                optimized.append(chunk)

            i += 1

        # بازنویسی ID‌ها
        for idx, chunk in enumerate(optimized):
            chunk['chunk_id'] = idx
            chunk['total_chunks'] = len(optimized)

        return optimized

    def _split_large_chunk(self, text: str, start_id: int) -> List[Dict]:
        """تقسیم chunk بزرگ"""

        sentences = self.sent_tokenizer.tokenize(text)
        sub_chunks = []
        current_chunk = []
        current_words = 0

        for sent in sentences:
            words = self.word_tokenizer.tokenize(sent)

            if current_words + len(words) > self.max_chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                sub_chunks.append(self._create_chunk_dict(chunk_text, start_id + len(sub_chunks), None))

                # overlap
                current_chunk = current_chunk[-2:] if len(current_chunk) >= 2 else []
                current_words = sum(len(self.word_tokenizer.tokenize(s)) for s in current_chunk)

            current_chunk.append(sent)
            current_words += len(words)

        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            sub_chunks.append(self._create_chunk_dict(chunk_text, start_id + len(sub_chunks), None))

        return sub_chunks

    def _extract_keywords(self, text: str, top_n: int = 5) -> List[str]:
        """استخراج کلمات کلیدی"""

        words = self.word_tokenizer.tokenize(text)
        stop_words = {'و', 'در', 'به', 'از', 'که', 'این', 'را', 'با', 'برای', 'آن', 'یک', 'است', 'شود', 'می', 'خود'}

        word_freq = {}
        for word in words:
            if len(word) > 2 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1

        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:top_n]]

    def _print_stats(self, total_chunks: int):
        """چاپ آمار نهایی"""

        print("\n" + "=" * 50)
        print("📊 آمار Chunking")
        print("=" * 50)
        print(f"  Chunk‌های نهایی: {total_chunks}")
        print(f"  مرزهای Stage 1 (Pattern): {self.stats['stage1_boundaries']}")
        print(f"  مرزهای Stage 2 (Coherence): {self.stats['stage2_boundaries']}")
        print(f"  مرزهای Stage 3 (Embedding): {self.stats['stage3_boundaries']}")
        print(f"  تعداد فراخوانی Embedding: {self.stats['embedding_calls']}")

        total_boundaries = sum([
            self.stats['stage1_boundaries'],
            self.stats['stage2_boundaries'],
            self.stats['stage3_boundaries']
        ])

        if total_boundaries > 0:
            stage3_percent = (self.stats['stage3_boundaries'] / total_boundaries) * 100
            print(f"  درصد استفاده از Embedding: {stage3_percent:.1f}%")

        print("=" * 50 + "\n")


def extract_headings_from_docx(doc_path: str) -> Tuple[str, List[Dict]]:
    """استخراج متن و heading‌ها از فایل Word"""

    doc = Document(doc_path)
    full_text_parts = []
    headings = []
    para_index = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            full_text_parts.append('')
            para_index += 1
            continue

        # بررسی سبک پاراگراف
        style_name = para.style.name.lower()

        # تشخیص heading
        if 'heading' in style_name:
            # استخراج سطح heading
            level = 1
            if 'heading 1' in style_name:
                level = 1
            elif 'heading 2' in style_name:
                level = 2
            elif 'heading 3' in style_name:
                level = 3

            headings.append({
                'text': text,
                'level': level,
                'para_index': para_index
            })

        full_text_parts.append(text)
        para_index += 1

    full_text = '\n'.join(full_text_parts)

    return full_text, headings


def process_word_file(input_path: str, output_path: str = None, ollama_model: str = "embeddinggemma"):
    """پردازش فایل Word با Optimized Hybrid Chunker"""

    print(f"📖 خواندن فایل: {input_path}")

    # استخراج متن و heading‌ها
    full_text, headings = extract_headings_from_docx(input_path)

    if not full_text.strip():
        print("❌ فایل خالی است!")
        return None, None

    print(f"✓ متن استخراج شد: {len(full_text)} کاراکتر")
    print(f"✓ {len(headings)} عنوان (Heading) یافت شد")

    # Chunking
    chunker = OptimizedHybridChunker(
        min_chunk_size=200,
        max_chunk_size=800,
        overlap_size=50,
        ollama_model=ollama_model
    )

    chunks = chunker.create_chunks(full_text, headings=headings)

    if not chunks:
        print("❌ هیچ chunk ایجاد نشد!")
        return None, None

    # ایجاد فایل خروجی
    if output_path is None:
        input_file = Path(input_path)
        output_path = input_file.parent / f"{input_file.stem}_chunked{input_file.suffix}"

    output_doc = Document()

    # نوشتن chunk‌ها
    for chunk in chunks:
        # تگ جداکننده
        tag = f"@@-{chunk['chunk_id']:04d}-@@"
        tag_para = output_doc.add_paragraph(tag)
        tag_para.runs[0].bold = True
        tag_para.runs[0].font.size = 140000  # 14pt

        # متن chunk
        output_doc.add_paragraph(chunk['text'])

        # metadata
        meta_parts = [
            f"کلمات: {chunk['word_count']}",
            f"جملات: {chunk['sentence_count']}",
            f"Stage: {chunk.get('boundary_stage', 'N/A')}",
            f"اطمینان: {chunk.get('boundary_confidence', 0):.2f}"
        ]

        if chunk.get('keywords'):
            meta_parts.append(f"کلیدواژه: {', '.join(chunk['keywords'])}")

        if chunk.get('heading_level'):
            meta_parts.append(f"Heading سطح {chunk['heading_level']}")

        meta_text = " | ".join(meta_parts)
        meta_para = output_doc.add_paragraph(meta_text)
        meta_para.runs[0].italic = True
        meta_para.runs[0].font.size = 100000  # 10pt

        # overlap (اختیاری - برای دیباگ)
        if chunk.get('overlap_next'):
            overlap_para = output_doc.add_paragraph(f"[Overlap: {chunk['overlap_next'][:100]}...]")
            overlap_para.runs[0].font.color.rgb = (150, 150, 150)
            overlap_para.runs[0].font.size = 90000  # 9pt

        # فاصله بین chunk‌ها
        output_doc.add_paragraph()
        output_doc.add_paragraph("─" * 80)
        output_doc.add_paragraph()

    # ذخیره فایل
    output_doc.save(output_path)

    print(f"\n✅ فایل خروجی ذخیره شد: {output_path}")

    # آمار نهایی
    total_words = sum(c['word_count'] for c in chunks)
    avg_words = total_words / len(chunks) if chunks else 0
    min_words = min(c['word_count'] for c in chunks) if chunks else 0
    max_words = max(c['word_count'] for c in chunks) if chunks else 0

    print("\n📈 آمار chunk‌ها:")
    print(f"  تعداد: {len(chunks)}")
    print(f"  میانگین کلمات: {avg_words:.0f}")
    print(f"  کمترین: {min_words} کلمات")
    print(f"  بیشترین: {max_words} کلمات")

    # توزیع بر اساس stage
    stage_dist = {}
    for chunk in chunks:
        stage = chunk.get('boundary_stage', 'unknown')
        stage_dist[stage] = stage_dist.get(stage, 0) + 1

    print("\n🎯 توزیع chunk‌ها بر اساس روش تشخیص:")
    stage_names = {0: 'Heading', 1: 'Pattern', 2: 'Coherence', 3: 'Embedding'}
    for stage, count in sorted(stage_dist.items()):
        stage_name = stage_names.get(stage, 'Unknown')
        print(f"  {stage_name}: {count} chunk")

    return chunks, str(output_path)


def test_ollama_connection(model: str = "embeddinggemma"):
    """تست اتصال به Ollama"""

    print("🔍 تست اتصال به Ollama...")

    embedder = OllamaEmbedding(model=model)
    test_text = "این یک متن تست است"

    result = embedder.get_embedding(test_text)

    if result is not None:
        print(f"✅ اتصال موفق! (embedding dimension: {len(result)})")
        return True
    else:
        print("❌ خطا در اتصال به Ollama!")
        print("\nلطفاً موارد زیر را بررسی کنید:")
        print("1. Ollama در حال اجرا است؟")
        print("   $ ollama serve")
        print(f"2. مدل {model} نصب شده است؟")
        print(f"   $ ollama pull {model}")
        print("3. Ollama روی پورت پیش‌فرض است؟ (11434)")
        return False


# ===========================================
# مثال استفاده
# ===========================================

if __name__ == "__main__":
    import sys

    # تست اتصال اولیه
    if not test_ollama_connection("embeddinggemma"):
        print("\n⚠️ بدون Ollama، فقط Stage 1 و 2 کار می‌کند (دقت کمتر)")
        response = input("ادامه می‌دهید؟ (y/n): ")
        if response.lower() != 'y':
            sys.exit(1)

    # مسیر فایل ورودی
    input_file = "your_book.docx"  # فایل خود را اینجا قرار دهید

    # یا از command line
    if len(sys.argv) > 1:
        input_file = sys.argv[1]

    try:
        chunks, output_file = process_word_file(
            input_file,
            ollama_model="embeddinggemma"
        )

        if chunks and output_file:
            print("\n" + "=" * 50)
            print("🎉 عملیات با موفقیت کامل شد!")
            print("=" * 50)

            # نمایش نمونه
            print("\n📝 نمونه اولین chunk:")
            print("-" * 50)
            first_chunk = chunks[0]
            print(f"ID: {first_chunk['chunk_id']}")
            print(f"تعداد کلمات: {first_chunk['word_count']}")
            print(f"روش تشخیص: Stage {first_chunk.get('boundary_stage', 'N/A')}")
            print(f"اطمینان: {first_chunk.get('boundary_confidence', 0):.2%}")
            print(f"\nمتن:\n{first_chunk['text'][:300]}...")

            if first_chunk.get('keywords'):
                print(f"\nکلمات کلیدی: {', '.join(first_chunk['keywords'])}")

            print("\n" + "=" * 50)
            print(f"📂 فایل نهایی: {output_file}")
            print("=" * 50)

    except FileNotFoundError:
        print(f"\n❌ خطا: فایل '{input_file}' یافت نشد!")
        print("\nراهنما:")
        print(f"  python {sys.argv[0]} path/to/your/file.docx")

    except Exception as e:
        print(f"\n❌ خطا: {str(e)}")
        import traceback

        traceback.print_exc()


# ===========================================
# توابع کمکی اضافی
# ===========================================

def analyze_chunks(chunks: List[Dict]):
    """تحلیل دقیق chunk‌ها"""

    if not chunks:
        print("هیچ chunk‌ای برای تحلیل وجود ندارد!")
        return

    print("\n" + "=" * 60)
    print("📊 تحلیل جامع Chunk‌ها")
    print("=" * 60)

    # آمار اولیه
    total = len(chunks)
    total_words = sum(c['word_count'] for c in chunks)

    print(f"\n1️⃣ آمار کلی:")
    print(f"   تعداد کل chunk‌ها: {total}")
    print(f"   تعداد کل کلمات: {total_words:,}")
    print(f"   میانگین کلمات در chunk: {total_words / total:.1f}")

    # توزیع طول
    print(f"\n2️⃣ توزیع طول chunk‌ها:")

    ranges = [
        (0, 150, "خیلی کوچک"),
        (150, 300, "کوچک"),
        (300, 400, "متوسط"),
        (400, 600, "بزرگ"),
        (600, float('inf'), "خیلی بزرگ")
    ]

    for min_w, max_w, label in ranges:
        count = sum(1 for c in chunks if min_w <= c['word_count'] < max_w)
        if count > 0:
            percentage = (count / total) * 100
            print(
                f"   {label:15} ({min_w:3}-{max_w if max_w != float('inf') else '∞':>3}): {count:3} chunk ({percentage:5.1f}%)")

    # توزیع بر اساس روش تشخیص
    print(f"\n3️⃣ روش تشخیص مرزها:")

    stage_names = {
        0: "📑 Heading (از ورد)",
        1: "⚡ Pattern (الگو)",
        2: "🔍 Coherence (انسجام)",
        3: "🧠 Embedding (معنایی)"
    }

    stage_counts = {}
    for chunk in chunks:
        stage = chunk.get('boundary_stage', 'unknown')
        stage_counts[stage] = stage_counts.get(stage, 0) + 1

    for stage in sorted(stage_counts.keys()):
        count = stage_counts[stage]
        percentage = (count / total) * 100
        name = stage_names.get(stage, "❓ نامشخص")
        print(f"   {name:30}: {count:3} ({percentage:5.1f}%)")

    # اطمینان مرزها
    print(f"\n4️⃣ توزیع اطمینان مرزها:")

    confidences = [c.get('boundary_confidence', 0) for c in chunks]
    high_conf = sum(1 for c in confidences if c >= 0.8)
    med_conf = sum(1 for c in confidences if 0.5 <= c < 0.8)
    low_conf = sum(1 for c in confidences if c < 0.5)

    print(f"   بالا (≥0.8):   {high_conf:3} ({high_conf / total * 100:5.1f}%)")
    print(f"   متوسط (0.5-0.8): {med_conf:3} ({med_conf / total * 100:5.1f}%)")
    print(f"   پایین (<0.5):  {low_conf:3} ({low_conf / total * 100:5.1f}%)")

    # کلمات کلیدی پرتکرار
    print(f"\n5️⃣ کلمات کلیدی پرتکرار در کل متن:")

    all_keywords = []
    for chunk in chunks:
        all_keywords.extend(chunk.get('keywords', []))

    keyword_freq = Counter(all_keywords)
    top_keywords = keyword_freq.most_common(10)

    for keyword, count in top_keywords:
        print(f"   {keyword:20}: {count} بار")

    print("\n" + "=" * 60 + "\n")


def export_chunks_to_json(chunks: List[Dict], output_path: str):
    """خروجی JSON برای embedding بعدی"""

    import json

    output_data = {
        'metadata': {
            'total_chunks': len(chunks),
            'total_words': sum(c['word_count'] for c in chunks),
            'avg_words_per_chunk': sum(c['word_count'] for c in chunks) / len(chunks) if chunks else 0
        },
        'chunks': chunks
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)

    print(f"✅ JSON ذخیره شد: {output_path}")